#!C:\Python27\python.exe

from docx import Document

doc = Document()
doc.add_paragraph()
print len(doc.paragraphs)
