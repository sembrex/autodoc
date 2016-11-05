#!C:\Python27\python.exe

import sys, os

from PyQt4 import QtCore, QtGui

from docx import Document
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from StringIO import StringIO

import dokumentasi_ui

Ui_Main = dokumentasi_ui.Ui_MainWindow

def main():
	app = QtGui.QApplication(sys.argv)
	window = Dokumentasi()
	window.show()
	sys.exit(app.exec_())

class Dokumentasi(QtGui.QMainWindow, Ui_Main):
	def __init__(self):
		super(Dokumentasi, self).__init__()
		self.setupUi(self)
		self.prepareUi()

		self.btnInput.clicked.connect(self.browseInput)
		self.btnOutput.clicked.connect(self.browseOutput)
		self.btnExit.clicked.connect(self.close)
		self.btnStart.clicked.connect(self.start)
		self.btnLog.clicked.connect(self.toggleLog)


	def prepareUi(self):
		self.setWindowTitle('AutoDoc v1.1 - Karogis')
		self.mTop.setValidator(QtGui.QDoubleValidator(0, 10, 2, self))
		self.mBottom.setValidator(QtGui.QDoubleValidator(0, 10, 2, self))
		self.mLeft.setValidator(QtGui.QDoubleValidator(0, 10, 2, self))
		self.mRight.setValidator(QtGui.QDoubleValidator(0, 10, 2, self))

		self.logger('====================================')
		self.logger('||          *****            www.karogis.com            *****         ||')
		self.logger('====================================\n\n')

		self.logVisible = False
		self.toggleLog()
		self.statusbar.showMessage('Siap')


	def browseInput(self):
		directory = QtGui.QFileDialog.getExistingDirectory(self, 'Pilih Folder')
		if directory:
			self.inputFolder.setText(directory)
			self.logger('Folder Sumber: ' + directory)

			self.dirs = 0
			self.images = 0
			self.evaluateDirectory(directory)

			self.logger('Jumlah Folder: ' + str(self.dirs))
			self.logger('Jumlah Foto: ' + str(self.images))
			self.statusbar.showMessage('Jumlah Folder: ' + str(self.dirs) + ', Jumlah Foto: ' + str(self.images))


	def evaluateDirectory(self, directory, strip = '-'):
		directory = str(directory)

		for file in os.listdir(directory):
			if file.lower().endswith('.jpg') or file.lower().endswith('.jpeg') or file.lower().endswith('.png'):
				self.logger(strip + ' Foto: ' + file)
				self.images += 1

		QtGui.QApplication.processEvents()

		for file in os.listdir(directory):
			if os.path.isdir(os.path.join(directory, file)):
				self.dirs += 1
				self.logger('=====================================')
				self.logger(strip + ' Folder: ' + file)
				self.evaluateDirectory(os.path.join(directory, file), strip + ' -')
				self.logger('=====================================')

		QtGui.QApplication.processEvents()


	def browseOutput(self):
		file = QtGui.QFileDialog.getSaveFileName(self, 'Simpan Sebagai', '', 'Word Document (*.docx)')
		if file:
			self.outputFolder.setText(file)
			self.logger('File Hasil: ' + file)


	def start(self):
		if self.inputFolder.text() == '' or self.outputFolder.text() == '':
			self.showDialog('error', 'Kesalahan!', 'Folder Sumber & File Hasil harus diisi')
			return
		elif self.images == 0:
			self.showDialog('error', 'Kesalahan!', 'Tidak ditemukan foto')
			return

		top_margin = float(self.mTop.text()) if self.mTop.text() != '' else 2.0
		bottom_margin = float(self.mBottom.text()) if self.mBottom.text() != '' else 1.5
		left_margin = float(self.mLeft.text()) if self.mLeft.text() != '' else 3.0
		right_margin = float(self.mRight.text()) if self.mRight.text() != '' else 1.5

		self.logger('\nMenggunakan pengaturan jarak: Atas = ' + str(top_margin) +
						'cm; Bawah = ' + str(bottom_margin) +
						'cm; Kiri = ' + str(left_margin) +
						'cm; Kanan = ' + str(right_margin) + 'cm')

		self.doc = Document()

		try:
			self.doc.save(str(self.outputFolder.text()))
		except Exception as e:
			self.logger('KESALAHAN: ' + e.strerror)
			self.showDialog('error', 'Kesalahan', e.strerror)
			return

		QtGui.QApplication.setOverrideCursor(QtGui.QCursor(QtCore.Qt.WaitCursor))
		self.logger('\n\nMemulai Proses...')
		self.logger('Ngopi sek coy...\n\n')
		self.statusbar.showMessage('Ngopi sek coy...')
		self.btnStart.setText('Udud dulu...')
		self.btnStart.setEnabled(False)
		self.btnExit.setEnabled(False)
		self.progressBar.setMaximum(self.images)
		self.progressBar.setValue(0)
		self.progressCounter = 0

		QtGui.QApplication.processEvents()

		for section in self.doc.sections:
			section.top_margin = Cm(top_margin)
			section.bottom_margin = Cm(bottom_margin)
			section.left_margin = Cm(left_margin)
			section.right_margin = Cm(right_margin)
			section.page_width = Inches(8.267)
			section.page_height = Inches(11.692)

		directory = str(self.inputFolder.text())

		self.insertPictures(directory)

		QtGui.QApplication.processEvents()

		QtGui.QApplication.restoreOverrideCursor()

		try:
			self.doc.save(str(self.outputFolder.text()))
		except Exception as e:
			self.logger('KESALAHAN: ' + e.strerror)
			self.statusbar.showMessage('Siap')
			self.btnStart.setText('Mulai')
			self.btnStart.setEnabled(True)
			self.btnExit.setEnabled(True)
			self.showDialog('error', 'Kesalahan', e.strerror)
		else:
			self.logger('Selesai!')
			self.statusbar.showMessage('Siap')
			self.btnStart.setText('Mulai')
			self.btnStart.setEnabled(True)
			self.btnExit.setEnabled(True)
			self.showDialog('info', 'Selesai coy...', 'Dokumentasi berhasil dibuat.\nBuka file?',
				[('Tidak', QtGui.QMessageBox.NoRole), ('Ya', QtGui.QMessageBox.YesRole)],
				self.openFile)


	def insertPictures(self, directory):
		counter = 0

		titlePrefix = str(self.titlePrefix.text()) + ' ' if self.titlePrefix.text() != '' else ''

		if self.countPictures(directory) > 0:
			heading = self.doc.add_heading(titlePrefix + os.path.basename(directory).upper())
			heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
			heading_format = heading.paragraph_format
			heading_format.space_before = Pt(0)
			heading_format.space_after = Pt(24)

			table = self.doc.add_table(rows=1, cols=2)
			cells = table.rows[0].cells

			for file in os.listdir(directory):
				if (file.lower().endswith('.jpg') or file.lower().endswith('.jpeg') or file.lower().endswith('.png')) :
					if counter % 6 == 0 and counter != 0:
						self.doc.add_page_break()
						h = self.doc.add_heading(titlePrefix + os.path.basename(directory).upper())
						h.alignment = WD_ALIGN_PARAGRAPH.CENTER
						hf = h.paragraph_format
						hf.space_before = Pt(0)
						hf.space_after = Pt(24)
						table = self.doc.add_table(rows=1, cols=2)
						cells = table.rows[0].cells
					elif counter % 2 == 0 and counter != 0:
						cells = table.add_row().cells

					self.logger('Memproses foto ' + file)
					self.progressCounter += 1
					self.progressBar.setValue(self.progressCounter)
					QtGui.QApplication.processEvents()

					img = Image.open(os.path.join(directory, file))
					output = StringIO()
					width = 600
					width_percent = float(width) / float(img.size[0])
					height = int(img.size[1] * width_percent)
					image = img.resize((width, height), Image.BICUBIC)
					image.save(output, 'JPEG', quality=80)

					run = cells[counter % 2].paragraphs[0].add_run()
					run.add_picture(output, width=Cm(8))

					cells[counter % 2].add_paragraph(os.path.splitext(file)[0]).alignment = WD_ALIGN_PARAGRAPH.CENTER
					counter += 1

			QtGui.QApplication.processEvents()

		for file in os.listdir(directory):
			if os.path.isdir(os.path.join(directory, file)):
				self.logger('=====================================')
				self.logger('Folder ditemukan ' + file)
				QtGui.QApplication.processEvents()
				if len(self.doc.paragraphs) > 0:
					self.doc.add_page_break()
				self.insertPictures(os.path.join(directory, file))

		QtGui.QApplication.processEvents()


	def countPictures(self, directory):
		counter = 0

		for file in os.listdir(directory):
			if (file.lower().endswith('.jpg') or file.lower().endswith('.jpeg') or file.lower().endswith('.png')) :
				counter += 1

		QtGui.QApplication.processEvents()
		return counter


	def openFile(self, reply):
		if reply == 1:
			try:
				os.startfile(str(self.outputFolder.text()))
				QtGui.QApplication.processEvents()
			except Exception as e:
				self.showDialog('error', 'Kesalahan', e.strerror)
				return
		return


	def toggleLog(self):
		if self.logVisible:
			self.label_3.hide()
			self.logs.hide()
			self.btnLog.setText('Tampilkan Log')
			self.logVisible = False
		else:
			self.label_3.show()
			self.logs.show()
			self.btnLog.setText('Sembunyikan Log')
			self.logVisible = True


	def logger(self, string):
		self.logs.insertPlainText(string + "\n")
		self.logs.ensureCursorVisible()


	def showDialog(self, type, title, text, buttons = None, callback = None):
		dialog = QtGui.QMessageBox()

		if type == 'info':
			icon = QtGui.QMessageBox.Information
		elif type == 'warning':
			icon = QtGui.QMessageBox.Warning
		elif type == 'error':
			icon = QtGui.QMessageBox.Critical
		elif type == 'question':
			icon = QtGui.QMessageBox.Question
		else:
			icon = QtGui.QMessageBox.Information

		ico = QtGui.QIcon()
		ico.addPixmap(QtGui.QPixmap(":/icon/folder.png"), QtGui.QIcon.Normal, QtGui.QIcon.Off)
		dialog.setIcon(icon)
		dialog.setWindowTitle(title)
		dialog.setWindowIcon(ico)
		dialog.setText(text)

		if buttons == None:
			dialog.setStandardButtons(QtGui.QMessageBox.Ok)
		else:
			for button, role in buttons:
				dialog.addButton(button, role)

		reply = dialog.exec_()

		if callable(callback):
			callback(reply)


	def closeEvent(self, event):
		reply = QtGui.QMessageBox.question(self, 'Keluar',
			"Sudah ngopi bro?", 'Belum', 'Sudah')

		if reply == 1:
			event.accept()
		else:
			event.ignore()
			self.showDialog('info', 'Ngopi', 'Ngopi dulu bro...')

if __name__ == '__main__':
	main()

